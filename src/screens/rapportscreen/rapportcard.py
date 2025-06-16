from typing import Any, List, Optional
from flet import *
from fpdf import FPDF
from docx import Document
import csv
import sqlite3

# import json
from myaction import recuperer_liste_rapports

from screens.rapportscreen.rapportupdateform import RapportUpdateForm
from uix.traitext import TraiText
# from uix.customdropdown import CustomDropDown

DB_PATH = "data/rapport.db"

class RapportCard(Card):
    def __init__(self, page: Page, rapport, formcontrol):
        super().__init__()
        self.expand=True
        self.elevation=1
        self.rapport=rapport
        self.formcontrol=formcontrol
        self.page=page
        self.content=Container(
            # on_click=lambda e: self.selectprojet(e),
            padding=padding.all(10),
            data=rapport,
            ink=True,
            expand=True,
            content=Container(
                content=Column([
                    Text(f"📝 {rapport['title']}"),
                    Text(f"Rapport du {rapport['rapport_date']}"),
                    Text(rapport['content'], size=12),
                    Text(f"Ajouté le {rapport['created_at']}", italic=True, size=10),
                    Row([
                        IconButton(icon=Icons.EDIT, on_click= self.show_edit_rapport),
                        IconButton(icon=Icons.DELETE, on_click= self.show_delete_rapport),
                        PopupMenuButton(items=[
                            PopupMenuItem(text="Générer PDF", on_click=self.showGenerate_pdf),
                            PopupMenuItem(text="Générer DOCX", on_click=self.generate_docx),
                            PopupMenuItem(text="Générer CSV", on_click=self.generate_csv),
                            ])
                        ],alignment=MainAxisAlignment.END
                    )
                ]),
                padding=10
            )
        )
        

    def show_edit_rapport(self,e):
        rapport_content = RapportUpdateForm(self.page, rapport=self.rapport, formcontrol= self)
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Modifier Rapport"),
            content=rapport_content,
            actions=[
                TextButton("Annuler", on_click=self.close_dlg),
                TextButton("Enregistrer", on_click=rapport_content.ConfirmeUpdate),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
            content_padding=0
        )
        self.page.open(self.dlg_modal)
        self.page.update()

    def show_delete_rapport(self,e):
        title=self.rapport['title']
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Confirmation"),
            content=Text(f"Voulez-vous supprimer {title} ?"),
            actions=[
                TextButton("Non", on_click=self.close_dlg),
                TextButton("Oui", on_click=self.del_rapport),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
            content_padding=0
        )
        self.page.open(self.dlg_modal)
        self.page.update()
        

    def generate_docx(self, e):
        report_id=self.rapport["rid"]
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT title, content FROM rapports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()
        if row:
            title, content = row
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
            doc.save(f"generated_docs/rapport_{report_id}.docx")
            print("succes generé")
            
    def showGenerate_pdf(self,e):
        title=self.rapport['title']
        titlefield=TextField(value=title, expand=True, height=40)
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Confirmation"),
            content=titlefield,
            actions=[
                TextButton("Annuler", on_click=self.close_dlg),
                TextButton("Exporter", on_click = lambda e : self.generate_pdf(titlefield.value)),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
        )
        self.page.open(self.dlg_modal)
        self.page.update()

    def generate_csv(self, e):
        report_id=self.rapport["rid"]
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT title, content, rapport_date FROM rapports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()
        if row:
            with open(f"generated_docs/rapport_{report_id}.csv", mode="w", newline="", encoding="utf-8") as file:
                writer = csv.writer(file)
                writer.writerow(["Titre", "Contenu", "Date"])
                writer.writerow(row)

    def generate_pdf(self, title):
        report_id=self.rapport["rid"]
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT title, content, rapport_date FROM rapports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()

        if row:
            title, content, rapport_date = row
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)

            # Titre du rapport
            pdf.cell(0, 10, title, ln=True, align='C')

            # Semaine
            pdf.set_font("Arial", "", 12)
            pdf.cell(0, 10, f"Rapport du {rapport_date}", ln=True, align='C')

            pdf.ln(10)

            # Contenu
            pdf.multi_cell(0, 10, content)

            # Sauvegarde
            file_path = f"generated_docs/rapport_{title}.pdf"
            pdf.output(file_path)

            self.page.open(SnackBar(Text(f"rapport_{title} est exporter avec succès")))
        self.close_dlg(e=None)

        
    def close_dlg(self,e):
        try:
            self.page.close(self.dlg_modal)
            self.page.update()
        except:
            pass
        
    def del_rapport(self,e):
        rid=int(self.rapport['rid'])
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM rapports WHERE id=?", (rid,))
        conn.commit()
        conn.close()
        self.page.close(self.dlg_modal)
        self.formcontrol.load_rapports()
        
    